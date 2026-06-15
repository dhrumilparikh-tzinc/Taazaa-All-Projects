import pytest

from app.processors.base import InvalidInputError, RateLimitError
from app.workers.tasks import classify_error

# ── classify_error tests ──────────────────────────────────────────────────────


def test_classify_error_rate_limit_429():
    err = Exception("429 quota exceeded")
    error_type, retryable = classify_error(err)
    assert error_type == "RATE_LIMIT"
    assert retryable is True


def test_classify_error_rate_limit_from_custom_exception():
    err = RateLimitError("429: Gemini rate limit")
    error_type, retryable = classify_error(err)
    assert error_type == "RATE_LIMIT"
    assert retryable is True


def test_classify_error_invalid_input_400():
    err = Exception("400 invalid argument")
    error_type, retryable = classify_error(err)
    assert error_type == "INVALID_INPUT"
    assert retryable is False


def test_classify_error_invalid_input_from_custom_exception():
    err = InvalidInputError("400: Gemini invalid argument")
    error_type, retryable = classify_error(err)
    assert error_type == "INVALID_INPUT"
    assert retryable is False


def test_classify_error_unknown_retryable():
    err = Exception("connection reset by peer")
    error_type, retryable = classify_error(err)
    assert error_type == "UNKNOWN"
    assert retryable is True


def test_classify_error_quota_string():
    err = Exception("quota limit exceeded")
    error_type, retryable = classify_error(err)
    assert error_type == "RATE_LIMIT"
    assert retryable is True


def test_classify_error_rate_string():
    err = Exception("rate limit hit")
    error_type, retryable = classify_error(err)
    assert error_type == "RATE_LIMIT"
    assert retryable is True


# ── PDF extraction tests ──────────────────────────────────────────────────────


def test_pdf_extraction_has_page_markers(tmp_path):
    """PDF extract() output must contain [Page N] markers."""
    from unittest.mock import MagicMock

    from fpdf import FPDF

    from app.processors.pdf import PDFProcessor

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "Test content page one")
    pdf.add_page()
    pdf.cell(0, 10, "Test content page two")
    pdf_path = str(tmp_path / "test.pdf")
    pdf.output(pdf_path)

    job = MagicMock()
    job.file_path = pdf_path
    job.file_type = "pdf"
    job.id = "test-id"
    settings = MagicMock()
    settings.GEMINI_API_KEY = "fake"

    processor = PDFProcessor.__new__(PDFProcessor)
    processor.job = job
    processor.settings = settings
    processor.log = MagicMock()

    text = processor.extract()
    assert "[Page 1]" in text
    assert "[Page 2]" in text
    assert len(text) > 0


def test_docx_extraction_returns_paragraphs(tmp_path):
    """DOCX extract() returns paragraph text."""
    from unittest.mock import MagicMock

    from docx import Document

    from app.processors.docx_proc import DOCXProcessor

    doc = Document()
    doc.add_paragraph("First paragraph about AI.")
    doc.add_paragraph("Second paragraph about RAG.")
    docx_path = str(tmp_path / "test.docx")
    doc.save(docx_path)

    job = MagicMock()
    job.file_path = docx_path
    job.file_type = "docx"
    job.id = "test-id"
    settings = MagicMock()
    settings.GEMINI_API_KEY = "fake"

    processor = DOCXProcessor.__new__(DOCXProcessor)
    processor.job = job
    processor.settings = settings
    processor.log = MagicMock()

    text = processor.extract()
    assert "First paragraph" in text
    assert "Second paragraph" in text
